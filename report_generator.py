from __future__ import annotations

import json
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as DocxImage
from docx.shared import Cm
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


IMAGE_MAX_HEIGHT = Cm(7)
ISSUE_DESCRIPTION_TOKEN = "<issue_description>"
ISSUE_IMAGES_TOKEN = "<issue_images>"
VERIFIER_NAME = "KHAIRUL ANUAR JOHARI"


@dataclass(slots=True)
class Issue:
    description: str
    image_paths: list[Path] = field(default_factory=list)


@dataclass(slots=True)
class ReportData:
    date: str
    project_name: str
    project_sub_name: str
    report_title: str
    report_purpose: str
    report_author: str
    report_author_role: str
    issues: list[Issue] = field(default_factory=list)

    @classmethod
    def from_dict(cls, payload: dict) -> "ReportData":
        issues = [
            Issue(
                description=item["description"].strip(),
                image_paths=[Path(path) for path in item.get("image_paths", [])],
            )
            for item in payload.get("issues", [])
        ]
        return cls(
            date=payload["date"].strip(),
            project_name=payload["project_name"].strip(),
            project_sub_name=payload["project_sub_name"].strip(),
            report_title=payload["report_title"].strip(),
            report_purpose=payload["report_purpose"].strip(),
            report_author=payload["report_author"].strip(),
            report_author_role=payload["report_author_role"].strip(),
            issues=issues,
        )

    def placeholder_map(self) -> dict[str, str]:
        return {
            "date": self.date,
            "project_name": self.project_name,
            "project_sub_name": self.project_sub_name,
            "report_title": self.report_title,
            "report_purpose": self.report_purpose,
            "report_author": self.report_author,
            "report_author_role": self.report_author_role,
        }


def render_report(
    template_path: str | Path,
    output_path: str | Path,
    report: ReportData,
) -> Path:
    document = Document(str(template_path))
    _populate_issue_table(document, report.issues)
    _replace_scalar_placeholders(document, report.placeholder_map())
    if report.report_author == VERIFIER_NAME:
        _remove_verifier_section(document)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def _replace_scalar_placeholders(document: DocxDocument, replacements: dict[str, str]) -> None:
    for paragraph in _iter_paragraphs(document):
        for name, value in replacements.items():
            _replace_placeholder_in_paragraph(paragraph, f"<{name}>", value)


def _iter_paragraphs(document: DocxDocument) -> Iterable[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _replace_placeholder_in_paragraph(paragraph: Paragraph, placeholder: str, value: str) -> None:
    if placeholder not in paragraph.text:
        return

    replaced_in_place = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            replaced_in_place = True

    if replaced_in_place:
        return

    if not paragraph.runs:
        paragraph.add_run(paragraph.text.replace(placeholder, value))
        return

    merged_text = "".join(run.text for run in paragraph.runs).replace(placeholder, value)
    paragraph.runs[0].text = merged_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _remove_verifier_section(document: DocxDocument) -> None:
    start_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if "Laporan Disahkan Oleh:" in paragraph.text:
            start_index = index
            break

    if start_index is None:
        return

    for paragraph in list(document.paragraphs[start_index:]):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _populate_issue_table(document: DocxDocument, issues: list[Issue]) -> None:
    table = _find_issue_table(document)
    if table is None:
        raise ValueError("Issue table placeholder was not found in the template.")

    issue_rows = table.rows[1:]
    if not issue_rows:
        raise ValueError("Issue table block is missing from the template.")

    template_block = [deepcopy(row._tr) for row in issue_rows]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    effective_issues = issues or [Issue(description="Tiada isu dalam pemantauan.", image_paths=[])]
    block_size = len(template_block)

    for index, issue in enumerate(effective_issues):
        for row_xml in deepcopy(template_block):
            table._tbl.append(row_xml)
        row_start = 1 + (index * block_size)
        _fill_issue_block(table, row_start, issue)


def _find_issue_table(document: DocxDocument) -> Table | None:
    for table in document.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if ISSUE_DESCRIPTION_TOKEN in text and ISSUE_IMAGES_TOKEN in text:
            return table
    return None


def _fill_issue_block(table: Table, row_start: int, issue: Issue) -> None:
    description_cell = table.cell(row_start, 0)
    image_cell = table.cell(row_start, 1)

    _set_cell_text(description_cell, issue.description)
    _set_issue_images(image_cell, issue.image_paths)


def _set_cell_text(cell: _Cell, text: str) -> None:
    cell.text = text


def _set_issue_images(cell: _Cell, image_paths: list[Path]) -> None:
    if not image_paths:
        cell.text = "Tiada lampiran."
        return

    cell.text = ""
    lines = _group_image_lines(image_paths)
    first_paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    for line_index, line in enumerate(lines):
        paragraph = first_paragraph if line_index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for image_index, image_path in enumerate(line):
            width, height = _scaled_dimensions(image_path)
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=width, height=height)
            if image_index < len(line) - 1:
                paragraph.add_run("  ")


def _group_image_lines(image_paths: list[Path]) -> list[list[Path]]:
    lines: list[list[Path]] = []
    pending_portraits: list[Path] = []

    for path in image_paths:
        if _is_portrait(path):
            pending_portraits.append(path)
            if len(pending_portraits) == 2:
                lines.append(pending_portraits)
                pending_portraits = []
            continue

        if pending_portraits:
            lines.append(pending_portraits)
            pending_portraits = []
        lines.append([path])

    if pending_portraits:
        lines.append(pending_portraits)

    return lines


def _is_portrait(image_path: Path) -> bool:
    image = DocxImage.from_file(str(image_path))
    return image.px_height > image.px_width


def _scaled_dimensions(image_path: Path) -> tuple[int, int]:
    image = DocxImage.from_file(str(image_path))
    native_width = int(image.width)
    native_height = int(image.height)
    max_height = int(IMAGE_MAX_HEIGHT)

    if native_height <= max_height:
        return native_width, native_height

    scale = max_height / native_height
    return int(native_width * scale), max_height


def load_report_data(payload_path: str | Path) -> ReportData:
    payload = json.loads(Path(payload_path).read_text(encoding="utf-8"))
    return ReportData.from_dict(payload)


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Render the initial report DOCX from JSON data.")
    parser.add_argument("payload", help="Path to a JSON file that matches ReportData.from_dict().")
    parser.add_argument(
        "--template",
        default="Template Initial Report.docx",
        help="Path to the DOCX template.",
    )
    parser.add_argument(
        "--output",
        default="output/Initial Report.docx",
        help="Where to save the rendered DOCX.",
    )
    args = parser.parse_args()

    report = load_report_data(args.payload)
    output = render_report(args.template, args.output, report)
    print(output)


if __name__ == "__main__":
    main()
