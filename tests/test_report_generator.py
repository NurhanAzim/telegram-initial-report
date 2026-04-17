from __future__ import annotations

import base64
import tempfile
import unittest
import zlib
import struct
from binascii import crc32
from pathlib import Path

from docx import Document
from docx.shared import Cm

from nextcloud_client import NextcloudClient, sanitize_filename_part
from report_generator import Issue, ReportData, render_report


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+jX3sAAAAASUVORK5CYII="
)


class ReportGeneratorTest(unittest.TestCase):
    def test_render_report_replaces_placeholders_and_adds_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "issue.png"
            image_path.write_bytes(PNG_1X1)

            template = Path("Template Initial Report.docx")
            output = temp_path / "rendered.docx"
            template_doc = Document(str(template))

            report = ReportData(
                date="16/04/2026",
                project_name="Projek Demo",
                project_sub_name="Fasa 1",
                report_title="Server Room",
                report_purpose="Pemeriksaan awal",
                report_author="MUHAMMAD ADAM BIN JAFFRY",
                report_author_role="DEVOPS ENGINEER",
                issues=[Issue(description="Kabel belum dirapikan", images_description="Lampiran utama", image_paths=[image_path])],
            )

            render_report(template, output, report)
            rendered_doc = Document(str(output))
            full_text = "\n".join(paragraph.text for paragraph in rendered_doc.paragraphs)
            table_text = "\n".join(cell.text for row in rendered_doc.tables[0].rows for cell in row.cells)

            self.assertIn("Tarikh Laporan: 16/04/2026", full_text)
            self.assertIn("Nama Projek: Projek Demo", full_text)
            self.assertIn("Sub-Projek: Fasa 1", full_text)
            self.assertIn("Tajuk: Laporan Server Room", full_text)
            self.assertIn("Tujuan: Pemeriksaan awal", full_text)
            self.assertIn("Nama: \tMUHAMMAD ADAM BIN JAFFRY", full_text)
            self.assertIn("Jawatan: DEVOPS ENGINEER", full_text)
            self.assertIn("Kabel belum dirapikan", table_text)
            self.assertIn("Lampiran utama", table_text)
            self.assertNotIn("<date>", full_text)
            self.assertNotIn("<issue_description>", table_text)
            self.assertNotIn("<issue_images>", table_text)
            self.assertGreater(rendered_doc.inline_shapes.__len__(), template_doc.inline_shapes.__len__())
            new_shape = rendered_doc.inline_shapes[-1]
            self.assertLessEqual(int(new_shape.height), int(Cm(7)))

    def test_render_report_expands_multiple_issue_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template = Path("Template Initial Report.docx")
            output = temp_path / "rendered-multi.docx"

            report = ReportData(
                date="16/04/2026",
                project_name="Projek Demo",
                project_sub_name="Fasa 1",
                report_title="Server Room",
                report_purpose="Pemeriksaan awal",
                report_author="MUHAMMAD ADAM BIN JAFFRY",
                report_author_role="DEVOPS ENGINEER",
                issues=[
                    Issue(description="Isu pertama", images_description="", image_paths=[]),
                    Issue(description="Isu kedua", images_description="", image_paths=[]),
                ],
            )

            render_report(template, output, report)
            rendered_doc = Document(str(output))
            table_text = "\n".join(cell.text for row in rendered_doc.tables[0].rows for cell in row.cells)

            self.assertIn("Isu pertama", table_text)
            self.assertIn("Isu kedua", table_text)
            self.assertEqual(len(rendered_doc.tables[0].rows), 7)

    def test_render_report_groups_portraits_two_per_line_and_landscape_one_per_line(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            portrait_1 = temp_path / "portrait-1.png"
            portrait_2 = temp_path / "portrait-2.png"
            landscape = temp_path / "landscape.png"
            _write_png(portrait_1, width=200, height=400)
            _write_png(portrait_2, width=220, height=440)
            _write_png(landscape, width=400, height=200)

            template = Path("Template Initial Report.docx")
            output = temp_path / "rendered-layout.docx"
            report = ReportData(
                date="16/04/2026",
                project_name="Projek Demo",
                project_sub_name="Fasa 1",
                report_title="Server Room",
                report_purpose="Pemeriksaan awal",
                report_author="MUHAMMAD ADAM BIN JAFFRY",
                report_author_role="DEVOPS ENGINEER",
                issues=[
                    Issue(
                        description="Susun atur imej",
                        images_description="Susunan atas gambar",
                        image_paths=[portrait_1, portrait_2, landscape],
                    )
                ],
            )

            render_report(template, output, report)
            rendered_doc = Document(str(output))
            image_cell = rendered_doc.tables[0].cell(1, 1)
            drawing_paragraphs = [
                paragraph
                for paragraph in image_cell.paragraphs
                if "pic:pic" in paragraph._element.xml
            ]

            self.assertEqual(image_cell.paragraphs[0].text, "Susunan atas gambar")
            self.assertEqual(len(drawing_paragraphs), 2)
            self.assertEqual(drawing_paragraphs[0]._element.xml.count("<pic:pic>"), 2)
            self.assertEqual(drawing_paragraphs[1]._element.xml.count("<pic:pic>"), 1)

    def test_render_report_hides_verifier_section_when_author_is_verifier(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template = Path("Template Initial Report.docx")
            output = temp_path / "rendered-verifier.docx"

            report = ReportData(
                date="16/04/2026",
                project_name="Projek Demo",
                project_sub_name="Fasa 1",
                report_title="Server Room",
                report_purpose="Pemeriksaan awal",
                report_author="KHAIRUL ANUAR JOHARI",
                report_author_role="TECHNICAL DIRECTOR",
                issues=[],
            )

            render_report(template, output, report)
            rendered_doc = Document(str(output))
            full_text = "\n".join(paragraph.text for paragraph in rendered_doc.paragraphs)

            self.assertIn("Nama: \tKHAIRUL ANUAR JOHARI", full_text)
            self.assertNotIn("Laporan Disahkan Oleh:", full_text)

    def test_nextcloud_share_url_parser(self) -> None:
        client = NextcloudClient(
            base_url="https://cloud.example.com",
            username="demo",
            password="secret",
            upload_dir="InitialReports",
        )
        body = """<?xml version="1.0"?>
<ocs>
  <meta>
    <status>ok</status>
    <statuscode>100</statuscode>
    <message>OK</message>
  </meta>
  <data>
    <url>https://cloud.example.com/s/demo123</url>
  </data>
</ocs>
"""
        share_id, share_url = client._extract_share_info(body)
        self.assertIsNone(share_id)
        self.assertEqual(share_url, "https://cloud.example.com/s/demo123")
        self.assertEqual(sanitize_filename_part("Projek Demo / Fasa 1"), "Projek-Demo-Fasa-1")

    def test_nextcloud_share_url_parser_accepts_namespaces_and_ok_status(self) -> None:
        client = NextcloudClient(
            base_url="https://cloud.example.com",
            username="demo",
            password="secret",
            upload_dir="InitialReports",
        )
        body = """<?xml version="1.0"?>
<ocs xmlns="http://open-collaboration-services.org/ns">
  <meta>
    <status>ok</status>
    <statuscode>200</statuscode>
    <message>OK</message>
  </meta>
  <data>
    <url>https://cloud.example.com/s/demo456</url>
  </data>
</ocs>
"""
        share_id, share_url = client._extract_share_info(body)
        self.assertIsNone(share_id)
        self.assertEqual(share_url, "https://cloud.example.com/s/demo456")


def _write_png(path: Path, width: int, height: int) -> None:
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", crc32(tag + data) & 0xFFFFFFFF)
        )

    header = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    row = b"\x00" + (b"\xff\xff\xff" * width)
    image_data = zlib.compress(row * height)
    path.write_bytes(
        header
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", image_data)
        + chunk(b"IEND", b"")
    )


if __name__ == "__main__":
    unittest.main()
